"""
Generate the Job Placement Agent Blueprint Word Document.
Run: python generate_blueprint.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour (hex without #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_para_bg(para, hex_color: str):
    """Set paragraph shading (used for code blocks)."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def add_horizontal_rule(doc):
    """Add a thin horizontal line."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2563EB")
    pBdr.append(bottom)
    pPr.append(pBdr)


def heading(doc, text, level=1, color="1E3A5F"):
    """Add a styled heading with colour."""
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return h


def body(doc, text, bold=False, italic=False, color=None, size=11):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_after = Pt(4)
    return p


def bullet(doc, text, level=0, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix followed by normal text."""
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    return p


def code_block(doc, lines):
    """Add a grey-background monospace code block."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.3)
        set_para_bg(p, "F1F5F9")
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("1E293B")


def info_box(doc, text, bg="EFF6FF", border_color="2563EB"):
    """Add a coloured info box as a 1x1 table."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("1E40AF")
    doc.add_paragraph()  # spacer


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Default body font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

# ===========================================================================
# COVER PAGE
# ===========================================================================
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run("JOB PLACEMENT AGENT")
title_run.bold = True
title_run.font.size = Pt(28)
title_run.font.color.rgb = RGBColor.from_string("1E3A5F")

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle_para.add_run("Project Blueprint & Architecture Design Document")
sub_run.font.size = Pt(16)
sub_run.font.color.rgb = RGBColor.from_string("2563EB")
sub_run.italic = True

doc.add_paragraph()

badge_para = doc.add_paragraph()
badge_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
badge_run = badge_para.add_run("Sprint 11  |  Milestone 1  |  Capstone Project")
badge_run.bold = True
badge_run.font.size = Pt(13)
badge_run.font.color.rgb = RGBColor.from_string("374151")

add_horizontal_rule(doc)

doc.add_paragraph()

meta_tbl = doc.add_table(rows=5, cols=2)
meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
labels = ["Course", "Sprint", "Milestone", "Author", "Date"]
values = [
    "Build and Deploy Intelligent Conversational Agents",
    "Sprint 11 — Capstone Project",
    "Milestone 1 — Core Feature Implementation",
    "Arockia Dhanaraj",
    "April 2026",
]
for i, (lbl, val) in enumerate(zip(labels, values)):
    row = meta_tbl.rows[i]
    set_cell_bg(row.cells[0], "1E3A5F")
    set_cell_bg(row.cells[1], "F8FAFC")
    lc = row.cells[0].paragraphs[0]
    lc.paragraph_format.left_indent = Pt(8)
    lr = lc.add_run(lbl)
    lr.bold = True
    lr.font.color.rgb = RGBColor(255, 255, 255)
    lr.font.size = Pt(11)
    vc = row.cells[1].paragraphs[0]
    vc.paragraph_format.left_indent = Pt(8)
    vr = vc.add_run(val)
    vr.font.size = Pt(11)
    vr.font.color.rgb = RGBColor.from_string("1E293B")

doc.add_page_break()

# ===========================================================================
# TABLE OF CONTENTS (manual)
# ===========================================================================
heading(doc, "Table of Contents", level=1)
toc_items = [
    ("1", "Project Overview & Objectives"),
    ("2", "Step 1 — Understand the Agent Workflow"),
    ("   2.1", "Agent Capabilities & Reasoning Flow"),
    ("   2.2", "Architecture Mapping"),
    ("3", "Step 2 — Get Started"),
    ("   3.1", "Technology Stack & Environment Setup"),
    ("   3.2", "Project Structure"),
    ("   3.3", "Pydantic Models & Schemas"),
    ("4", "Step 3 — Build Core Agent Features"),
    ("   4.1", "FastAPI Endpoints"),
    ("   4.2", "LLM Configuration (Gemini via LangChain)"),
    ("   4.3", "Multi-step LCEL Workflow"),
    ("   4.4", "Authentication (Auth0)"),
    ("   4.5", "Structured Validation"),
    ("   4.6", "Streaming Responses"),
    ("   4.7", "Langfuse Observability Tracing"),
    ("5", "Full System Architecture Diagram"),
    ("6", "Data Flow — End-to-End Example"),
    ("7", "Error Handling Strategy"),
    ("8", "Testing Strategy"),
    ("9", "Deployment Configuration"),
    ("10", "Environment Variables Reference"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{num}   {title}")
    r.font.size = Pt(11)
    if not num.startswith(" "):
        r.bold = True

doc.add_page_break()

# ===========================================================================
# SECTION 1 — PROJECT OVERVIEW
# ===========================================================================
heading(doc, "1. Project Overview & Objectives", level=1)
add_horizontal_rule(doc)

body(doc,
     "The Job Placement Agent is an AI-powered career assistant built as a full-stack "
     "conversational application. It helps job seekers discover opportunities, analyse "
     "their resume, identify skill gaps, and generate personalised cover letters — all "
     "through a natural-language chat interface.",
     size=11)

doc.add_paragraph()
heading(doc, "Core Capabilities", level=2, color="2563EB")

capabilities = [
    ("Job Discovery", "Real-time Google Jobs search via SerpAPI — returns up to 10 formatted listings with apply links."),
    ("Resume Analysis", "LLM-powered extraction of skills, experience level, education, and a quality score (out of 10)."),
    ("Skill Gap Analysis", "Compares a resume against a job description to surface matched skills, critical gaps, and a match percentage."),
    ("Cover Letter Generation", "Produces a 300–400-word, 4-paragraph personalised cover letter tailored to a specific role and company."),
]
for cap, desc in capabilities:
    bullet(doc, f" — {desc}", bold_prefix=cap)

doc.add_paragraph()
heading(doc, "Cross-cutting Features", level=2, color="2563EB")
cross = [
    "Conversational interface with per-session memory (LangChain + ChatMessageHistory)",
    "Auth0 RS256 JWT authentication on protected endpoints",
    "Uniform ApiResponse[T] envelope across all endpoints",
    "Langfuse observability — traces every LLM call, tool invocation, latency, and error",
    "Graceful degradation — Langfuse and Auth0 failures do not crash the application",
    "Vercel serverless deployment via Mangum ASGI adapter",
]
for item in cross:
    bullet(doc, item)

doc.add_page_break()

# ===========================================================================
# SECTION 2 — STEP 1: UNDERSTAND THE AGENT WORKFLOW
# ===========================================================================
heading(doc, "2. Step 1 — Understand the Agent Workflow", level=1)
add_horizontal_rule(doc)
info_box(doc,
         "Goal: Identify what the agent must do — prompts, tools, memory, and reasoning "
         "flow. Map architecture across LLM configs, chains, tool calls, auth, and UI.")

# 2.1
heading(doc, "2.1  Agent Capabilities & Reasoning Flow", level=2, color="2563EB")
body(doc, "The agent follows a structured reasoning loop built on LangChain's tool-calling "
         "agent pattern:")

doc.add_paragraph()
steps = [
    ("Receive", "User message arrives with a session_id for memory lookup."),
    ("Plan",    "The LLM decides which tool(s) to invoke based on the user intent."),
    ("Execute", "Tool is called (SerpAPI, resume LLM, or cover-letter LLM)."),
    ("Observe", "Tool result is returned to the agent as an observation."),
    ("Respond", "Agent synthesises observations into a final markdown-formatted reply."),
    ("Persist", "Exchange is saved to ChatMessageHistory for the session."),
]
flow_tbl = doc.add_table(rows=len(steps) + 1, cols=3)
flow_tbl.style = "Table Grid"
headers = ["#", "Phase", "Description"]
hdr_colors = ["1E3A5F", "1E3A5F", "1E3A5F"]
for i, (hdr, col) in enumerate(zip(headers, hdr_colors)):
    c = flow_tbl.rows[0].cells[i]
    set_cell_bg(c, col)
    p = c.paragraphs[0]
    r = p.add_run(hdr)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size = Pt(11)

for idx, (phase, desc) in enumerate(steps, start=1):
    row = flow_tbl.rows[idx]
    bg = "EFF6FF" if idx % 2 == 0 else "FFFFFF"
    for c in row.cells:
        set_cell_bg(c, bg)
    row.cells[0].paragraphs[0].add_run(str(idx)).font.size = Pt(11)
    pr = row.cells[1].paragraphs[0].add_run(phase)
    pr.bold = True
    pr.font.size = Pt(11)
    pr.font.color.rgb = RGBColor.from_string("1E3A5F")
    row.cells[2].paragraphs[0].add_run(desc).font.size = Pt(11)

doc.add_paragraph()

# 2.2
heading(doc, "2.2  Architecture Mapping", level=2, color="2563EB")

arch_items = [
    ("LLM Config",    "Google Gemini (gemini-2.0-flash) accessed through an OpenAI-compatible endpoint via LangChain's ChatOpenAI wrapper."),
    ("Chains",        "LCEL RunnableWithMessageHistory wraps the AgentExecutor for stateful, per-session conversation chains."),
    ("Tool Calls",    "Three LangChain @tool-decorated functions: search_jobs, analyze_resume, generate_cover_letter."),
    ("Auth",          "Auth0 RS256 JWT validated by python-jose; public endpoint bypasses auth for demo use."),
    ("UI",            "Streamlit frontend communicates with the FastAPI backend over HTTP REST."),
    ("Observability", "Langfuse CallbackHandler injected into AgentExecutor for automatic trace capture."),
]
arch_tbl = doc.add_table(rows=len(arch_items) + 1, cols=2)
arch_tbl.style = "Table Grid"
set_cell_bg(arch_tbl.rows[0].cells[0], "1E3A5F")
set_cell_bg(arch_tbl.rows[0].cells[1], "1E3A5F")
for cell, hdr in zip(arch_tbl.rows[0].cells, ["Component", "Details"]):
    r = cell.paragraphs[0].add_run(hdr)
    r.bold = True; r.font.color.rgb = RGBColor(255, 255, 255); r.font.size = Pt(11)

for idx, (comp, detail) in enumerate(arch_items, start=1):
    row = arch_tbl.rows[idx]
    bg = "EFF6FF" if idx % 2 == 0 else "FFFFFF"
    set_cell_bg(row.cells[0], bg); set_cell_bg(row.cells[1], bg)
    cr = row.cells[0].paragraphs[0].add_run(comp)
    cr.bold = True; cr.font.size = Pt(11); cr.font.color.rgb = RGBColor.from_string("1E3A5F")
    row.cells[1].paragraphs[0].add_run(detail).font.size = Pt(11)

doc.add_page_break()

# ===========================================================================
# SECTION 3 — STEP 2: GET STARTED
# ===========================================================================
heading(doc, "3. Step 2 — Get Started", level=1)
add_horizontal_rule(doc)
info_box(doc,
         "Goal: Set up FastAPI, LangChain, and the project environment. Create prompts, "
         "chains, tools, services, memory, and routes. Define core Pydantic models.")

# 3.1
heading(doc, "3.1  Technology Stack & Environment Setup", level=2, color="2563EB")

stack = [
    ("FastAPI ≥0.110",         "Web framework — async REST API with automatic OpenAPI docs"),
    ("Uvicorn ≥0.29",          "ASGI server for running FastAPI locally and in production"),
    ("LangChain 0.2.x",        "Agent orchestration, LCEL chains, tool calling, and memory"),
    ("LangChain-OpenAI ≥0.1",  "ChatOpenAI wrapper used to call the Gemini endpoint"),
    ("Google Gemini API",       "Primary LLM — gemini-2.0-flash via OpenAI-compatible endpoint"),
    ("SerpAPI",                 "Real-time Google Jobs search results"),
    ("Auth0",                   "RS256 JWT authentication and token management"),
    ("Langfuse ≥2.0",          "LLM observability, tracing, and cost monitoring"),
    ("Pydantic ≥2.0",          "Request/response validation and serialisation"),
    ("Streamlit ≥1.35",        "Chat UI — sidebar for user info, main area for conversation"),
    ("python-jose[cryptography]", "RS256 JWT validation"),
    ("Mangum ≥0.17",           "ASGI adapter for Vercel / AWS Lambda serverless deployment"),
]
st_tbl = doc.add_table(rows=len(stack) + 1, cols=2)
st_tbl.style = "Table Grid"
for i, (hdr, bg) in enumerate(zip(["Package / Service", "Purpose"], ["1E3A5F", "1E3A5F"])):
    c = st_tbl.rows[0].cells[i]
    set_cell_bg(c, bg)
    r = c.paragraphs[0].add_run(hdr)
    r.bold = True; r.font.color.rgb = RGBColor(255, 255, 255); r.font.size = Pt(11)

for idx, (pkg, purpose) in enumerate(stack, start=1):
    row = st_tbl.rows[idx]
    bg = "EFF6FF" if idx % 2 == 0 else "FFFFFF"
    set_cell_bg(row.cells[0], bg); set_cell_bg(row.cells[1], bg)
    cr = row.cells[0].paragraphs[0].add_run(pkg)
    cr.font.name = "Courier New"; cr.font.size = Pt(10); cr.font.color.rgb = RGBColor.from_string("7C3AED")
    row.cells[1].paragraphs[0].add_run(purpose).font.size = Pt(11)

doc.add_paragraph()

# 3.2
heading(doc, "3.2  Project Structure", level=2, color="2563EB")
code_block(doc, [
    "course2-capstoneproject/",
    "├── src/",
    "│   ├── backend/                   # FastAPI application",
    "│   │   ├── main.py                # App entry point, all endpoints",
    "│   │   ├── exceptions.py          # Domain exception hierarchy",
    "│   │   ├── requirements.txt",
    "│   │   ├── vercel.json",
    "│   │   ├── agent/",
    "│   │   │   ├── job_agent.py       # LangChain AgentExecutor + session memory",
    "│   │   │   ├── llm.py             # Gemini LLM factory & error translation",
    "│   │   │   ├── prompts.py         # System prompt template",
    "│   │   │   └── tools/",
    "│   │   │       ├── job_search.py       # SerpAPI Google Jobs tool",
    "│   │   │       ├── resume_analyzer.py  # LLM resume analysis tool",
    "│   │   │       └── cover_letter.py     # LLM cover letter generation tool",
    "│   │   ├── auth/",
    "│   │   │   └── auth0.py           # Auth0 RS256 JWT validation",
    "│   │   ├── models/",
    "│   │   │   ├── schemas.py         # Pydantic request/response models",
    "│   │   │   └── responses.py       # Uniform ApiResponse[T] envelope",
    "│   │   └── observability/",
    "│   │       └── langfuse_config.py # Langfuse callback handler factory",
    "│   └── frontend/",
    "│       ├── app.py                 # Streamlit chat UI",
    "│       └── requirements.txt",
    "├── tests/                         # 11 test modules (pytest)",
    "└── requirements-test.txt",
])

doc.add_paragraph()

# 3.3
heading(doc, "3.3  Pydantic Models & Schemas", level=2, color="2563EB")
body(doc, "All request and response shapes are defined in src/backend/models/schemas.py "
         "with strict validation rules:")

schema_groups = [
    ("Auth",         ["TokenRequest — username (email), password", "TokenResponse — access_token, token_type, expires_in"]),
    ("Chat",         ["ChatRequest — message (1–4000 chars), session_id (optional)", "ChatResponse — response text, session_id"]),
    ("Jobs",         ["JobSearchRequest — query (2–500 chars), location (optional)", "JobListing — title, company, location, qualifications, apply_link", "JobSearchResponse — list of JobListing, query, location, count"]),
    ("Resume",       ["ResumeAnalysisRequest — resume_text (50–50 000 chars), job_description (optional)", "ResumeAnalysisResponse — analysis text"]),
    ("Cover Letter", ["CoverLetterRequest — resume_text, job_description, job_title, company_name", "CoverLetterResponse — cover_letter text"]),
    ("System",       ["HealthResponse, SessionClearResponse, SessionCreateResponse, SessionsListResponse"]),
]
for group, items in schema_groups:
    p = doc.add_paragraph()
    r = p.add_run(group)
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor.from_string("1E3A5F")
    for item in items:
        bullet(doc, item, level=1)

doc.add_paragraph()
body(doc, "The ApiResponse[T] generic envelope (models/responses.py) wraps every endpoint:", bold=True)
code_block(doc, [
    '{',
    '  "success": true,',
    '  "data":    { ...typed payload... },',
    '  "error":   null,',
    '  "meta":    { "request_id": "<uuid>", "api_version": "1.0.0" }',
    '}',
])

doc.add_page_break()

# ===========================================================================
# SECTION 4 — STEP 3: BUILD CORE AGENT FEATURES
# ===========================================================================
heading(doc, "4. Step 3 — Build Core Agent Features", level=1)
add_horizontal_rule(doc)
info_box(doc,
         "Goal: Implement endpoints that run chains, tools, and memory logic. Configure "
         "LLMs and prompt templates. Build multi-step LCEL workflows. Add auth, structured "
         "validation, streaming responses, and Langfuse tracing.")

# 4.1
heading(doc, "4.1  FastAPI Endpoints", level=2, color="2563EB")
endpoints = [
    ("GET",    "/api/health",                    "Public",        "Liveness check — returns version, timestamp"),
    ("POST",   "/api/auth/token",                "Public",        "Auth0 Resource Owner Password Grant — returns JWT"),
    ("POST",   "/api/chat",                      "Auth Required", "Authenticated chat — runs agent with user context"),
    ("POST",   "/api/chat/public",               "Public",        "Demo chat — no token required; same agent logic"),
    ("POST",   "/api/session",                   "Public",        "Create a new session ID"),
    ("DELETE", "/api/chat/{session_id}",         "Public",        "Clear conversation history for a session"),
    ("GET",    "/api/sessions",                  "Public",        "List active in-memory sessions"),
    ("POST",   "/api/jobs/search",               "Auth Required", "Direct job search bypassing agent"),
    ("POST",   "/api/resume/analyze",            "Auth Required", "Direct resume analysis bypassing agent"),
    ("POST",   "/api/cover-letter/generate",     "Auth Required", "Direct cover letter generation"),
]
ep_tbl = doc.add_table(rows=len(endpoints) + 1, cols=4)
ep_tbl.style = "Table Grid"
for i, hdr in enumerate(["Method", "Path", "Auth", "Description"]):
    c = ep_tbl.rows[0].cells[i]
    set_cell_bg(c, "1E3A5F")
    r = c.paragraphs[0].add_run(hdr)
    r.bold = True; r.font.color.rgb = RGBColor(255, 255, 255); r.font.size = Pt(10)

method_colors = {"GET": "DCFCE7", "POST": "DBEAFE", "DELETE": "FEE2E2"}
for idx, (method, path, auth, desc) in enumerate(endpoints, start=1):
    row = ep_tbl.rows[idx]
    bg = "F8FAFC" if idx % 2 == 0 else "FFFFFF"
    set_cell_bg(row.cells[0], method_colors.get(method, bg))
    for ci in [1, 2, 3]:
        set_cell_bg(row.cells[ci], bg)
    mr = row.cells[0].paragraphs[0].add_run(method)
    mr.bold = True; mr.font.size = Pt(10); mr.font.color.rgb = RGBColor.from_string("1E3A5F")
    pr = row.cells[1].paragraphs[0].add_run(path)
    pr.font.name = "Courier New"; pr.font.size = Pt(9)
    ar = row.cells[2].paragraphs[0].add_run(auth)
    ar.font.size = Pt(10)
    if auth == "Auth Required":
        ar.font.color.rgb = RGBColor.from_string("B45309")
    dr = row.cells[3].paragraphs[0].add_run(desc)
    dr.font.size = Pt(10)

doc.add_paragraph()

# 4.2
heading(doc, "4.2  LLM Configuration (Gemini via LangChain)", level=2, color="2563EB")
body(doc, "The LLM is a Google Gemini model accessed through an OpenAI-compatible REST "
         "endpoint, allowing use of LangChain's ChatOpenAI class without a separate "
         "Google SDK:")
code_block(doc, [
    "# agent/llm.py",
    "from langchain_openai import ChatOpenAI",
    "",
    "def get_llm() -> ChatOpenAI:          # @lru_cache singleton",
    "    return ChatOpenAI(",
    "        model=settings.GEMINI_MODEL_NAME,   # gemini-2.0-flash",
    "        openai_api_key=settings.GEMINI_API_KEY,",
    "        openai_api_base=settings.GEMINI_BASE_URL,",
    "        temperature=0.7,",
    "    )",
])
doc.add_paragraph()
body(doc, "Error Translation Layer — llm.py maps raw OpenAI SDK errors to domain-specific "
         "exceptions with correct HTTP status codes:", bold=False)
err_pairs = [
    ("AuthenticationError", "GeminiConfigError (HTTP 500)"),
    ("RateLimitError",      "GeminiRateLimitError (HTTP 429)"),
    ("APIStatusError 429",  "GeminiQuotaExceededError (HTTP 503)"),
    ("APIConnectionError",  "GeminiNetworkError (HTTP 503)"),
    ("BadRequestError",     "GeminiInvalidRequestError (HTTP 400)"),
]
for sdk_err, domain_err in err_pairs:
    bullet(doc, f"  →  {domain_err}", bold_prefix=sdk_err)

doc.add_paragraph()

# 4.3
heading(doc, "4.3  Multi-step LCEL Workflow", level=2, color="2563EB")
body(doc, "The agent uses LangChain Expression Language (LCEL) to compose an agentic loop "
         "that handles multi-turn conversation and tool orchestration:")
code_block(doc, [
    "# agent/job_agent.py — simplified",
    "",
    "agent = create_tool_calling_agent(",
    "    llm=get_llm(),",
    "    tools=[search_jobs, analyze_resume, generate_cover_letter],",
    "    prompt=ChatPromptTemplate.from_messages([",
    '        ("system",        SYSTEM_PROMPT),',
    '        MessagesPlaceholder("chat_history"),',
    '        ("human",         "{input}"),',
    '        MessagesPlaceholder("agent_scratchpad"),',
    "    ]),",
    ")",
    "",
    "executor = AgentExecutor(agent=agent, tools=[...], verbose=False)",
    "",
    "chain = RunnableWithMessageHistory(",
    "    executor,",
    "    get_session_history,          # returns ChatMessageHistory per session",
    '    input_messages_key="input",',
    '    history_messages_key="chat_history",',
    ")",
    "",
    '# Invocation',
    'result = chain.invoke(',
    '    {"input": user_message},',
    '    config={"configurable": {"session_id": session_id}}',
    ')',
])

doc.add_paragraph()

# 4.4
heading(doc, "4.4  Authentication (Auth0)", level=2, color="2563EB")
body(doc, "Protected endpoints require a valid Auth0 RS256 JWT passed as a Bearer token:")
auth_steps = [
    "Client calls POST /api/auth/token with username + password.",
    "Backend calls Auth0 Resource Owner Password Grant → returns access_token.",
    "Client includes Authorization: Bearer <token> on subsequent requests.",
    "get_current_user() FastAPI dependency extracts and validates the JWT:",
    "     a. Fetches JWKS from Auth0 (cached in memory).",
    "     b. Validates RS256 signature, audience, issuer, and expiry.",
    "     c. Returns decoded payload or raises Auth0CredentialsError (401).",
    "Public /api/chat/public uses get_optional_user() — returns None if no token.",
]
for step in auth_steps:
    bullet(doc, step)

doc.add_paragraph()

# 4.5
heading(doc, "4.5  Structured Validation", level=2, color="2563EB")
body(doc, "Validation is layered across the stack:")
val_items = [
    ("Pydantic (schemas.py)",  "Field-level constraints, custom validators, extra='forbid'. Blank/whitespace rejection on all text fields."),
    ("FastAPI",                "Automatic 422 Unprocessable Entity on schema violations; mapped to ApiResponse envelope by global handler."),
    ("Tool layer",             "Each tool validates its inputs before invoking LLM or SerpAPI (min length checks, required fields)."),
    ("Exception hierarchy",    "exceptions.py provides typed domain errors carrying http_status and error_code for automatic handler routing."),
]
for label, desc in val_items:
    bullet(doc, f" — {desc}", bold_prefix=label)

doc.add_paragraph()

# 4.6
heading(doc, "4.6  Streaming Responses", level=2, color="2563EB")
body(doc, "The architecture is designed for streaming extensibility:")
stream_items = [
    "AgentExecutor supports streaming=True to yield partial tokens as they arrive.",
    "FastAPI's StreamingResponse with EventSourceResponse enables Server-Sent Events (SSE).",
    "The Streamlit frontend polls the backend or can be upgraded to consume SSE for real-time display.",
    "Current implementation returns complete responses; streaming can be enabled per-endpoint by switching to StreamingResponse and iterating chain.astream().",
]
for item in stream_items:
    bullet(doc, item)

doc.add_paragraph()

# 4.7
heading(doc, "4.7  Langfuse Observability Tracing", level=2, color="2563EB")
body(doc, "Every LLM call and tool invocation is traced automatically via Langfuse:")
code_block(doc, [
    "# observability/langfuse_config.py",
    "",
    "def get_langfuse_handler(session_id, user_id) -> CallbackHandler | None:",
    "    return CallbackHandler(",
    "        secret_key=LANGFUSE_SECRET_KEY,",
    "        public_key=LANGFUSE_PUBLIC_KEY,",
    "        host=LANGFUSE_HOST,",
    '        session_id=session_id,',
    '        user_id=user_id,',
    '        tags=["job-agent", "langchain"],',
    "    )",
    "",
    "# Injected into AgentExecutor via callbacks=[handler]",
    "# Gracefully returns None if credentials are missing",
])
doc.add_paragraph()
langfuse_items = [
    "Captures: LLM token counts, latency, tool call sequences, error traces.",
    "Session grouping allows full conversation replay in Langfuse dashboard.",
    "Graceful degradation — missing credentials → None returned → no handler injected.",
    "flush_langfuse() called on application shutdown to send pending events.",
]
for item in langfuse_items:
    bullet(doc, item)

doc.add_page_break()

# ===========================================================================
# SECTION 5 — ARCHITECTURE DIAGRAM  (visual table-based)
# ===========================================================================
heading(doc, "5. Full System Architecture Diagram", level=1)
add_horizontal_rule(doc)
body(doc, "The diagram below shows each layer of the system — from the user's browser "
         "down to external services — as colour-coded boxes connected by labelled arrows.")
doc.add_paragraph()

# ── helpers ────────────────────────────────────────────────────────────────

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set individual borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), str(val.get("sz", 12)))
            el.set(qn("w:color"), val.get("color", "000000"))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def arch_cell(cell, title, subtitle, pills,
              bg="1E3A5F", title_color="FFFFFF",
              subtitle_color="BFD4F2", pill_bg="FFFFFF", pill_fg="1E3A5F"):
    """Fill an architecture diagram cell with title, subtitle, and pill badges."""
    set_cell_bg(cell, bg)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # title
    p_title = cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(2)
    rt = p_title.add_run(title)
    rt.bold = True
    rt.font.size = Pt(12)
    rt.font.color.rgb = RGBColor.from_string(title_color)

    # subtitle
    if subtitle:
        p_sub = cell.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(4)
        rs = p_sub.add_run(subtitle)
        rs.italic = True
        rs.font.size = Pt(9)
        rs.font.color.rgb = RGBColor.from_string(subtitle_color)

    # pill badges (shown as a comma-separated inline row)
    if pills:
        p_pills = cell.add_paragraph()
        p_pills.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_pills.paragraph_format.space_before = Pt(2)
        p_pills.paragraph_format.space_after = Pt(6)
        for i, pill in enumerate(pills):
            if i:
                sep = p_pills.add_run("   ")
                sep.font.size = Pt(9)
            rp = p_pills.add_run(f"  {pill}  ")
            rp.font.size = Pt(8)
            rp.font.color.rgb = RGBColor.from_string(pill_fg)
            rp.bold = True


def arrow_row(tbl, row_idx, label="", ncols=3):
    """Merge a table row into one cell and show a down-arrow with an optional label."""
    row = tbl.rows[row_idx]
    merged = row.cells[0]
    for c in row.cells[1:]:
        merged = merged.merge(c)
    set_cell_bg(merged, "F8FAFC")
    p = merged.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    arrow_run = p.add_run("↓")
    arrow_run.font.size = Pt(16)
    arrow_run.font.color.rgb = RGBColor.from_string("2563EB")
    arrow_run.bold = True
    if label:
        lbl_run = p.add_run(f"  {label}")
        lbl_run.font.size = Pt(9)
        lbl_run.italic = True
        lbl_run.font.color.rgb = RGBColor.from_string("6B7280")


# ── diagram layout ──────────────────────────────────────────────────────────
# Row layout  (each row spans all 3 columns unless split)
# 0  : LAYER LABEL — User Interface
# 1  : Streamlit Frontend  [full width]
# 2  : arrow  "HTTP REST  /api/chat/public"
# 3  : LAYER LABEL — API Layer
# 4  : FastAPI Backend     [full width]
# 5  : Auth0 badge row     [full width, lighter]
# 6  : arrow  "run_agent()"
# 7  : LAYER LABEL — Agent / AI Layer
# 8  : LangChain Agent     [full width]
# 9  : Memory + Langfuse   [col 0+1 | col 2]
# 10 : arrow  "Tool calls"
# 11 : LAYER LABEL — Tools
# 12 : search_jobs | analyze_resume | generate_cover_letter  [3 cols]
# 13 : arrow  (3 cols)
# 14 : LAYER LABEL — External Services
# 15 : SerpAPI | Gemini (Analysis) | Gemini (Generation)  [3 cols]
# 16 : External services footer

NROWS = 17
NCOLS = 3

diag = doc.add_table(rows=NROWS, cols=NCOLS)
diag.alignment = WD_TABLE_ALIGNMENT.CENTER

# helper: merge entire row
def full_merge(row_idx):
    row = diag.rows[row_idx]
    merged = row.cells[0]
    for c in row.cells[1:]:
        merged = merged.merge(c)
    return merged

# helper: section label banner
def layer_label(row_idx, text, bg="374151"):
    cell = full_merge(row_idx)
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.color.rgb = RGBColor.from_string("D1D5DB")

# ── Row 0: UI layer label ──
layer_label(0, "User Interface Layer", bg="1D4ED8")

# ── Row 1: Streamlit Frontend ──
fe = full_merge(1)
arch_cell(fe,
          title="Streamlit Frontend",
          subtitle="src/frontend/app.py  ·  Port 8501",
          pills=["Chat UI", "Sidebar: User Info + Resume", "Session Management", "Quick Actions"],
          bg="2563EB", title_color="FFFFFF", subtitle_color="BFDBFE",
          pill_bg="1D4ED8", pill_fg="DBEAFE")

# ── Row 2: arrow ──
arrow_row(diag, 2, label="HTTP REST  →  POST /api/chat/public  |  GET /api/health  |  POST /api/auth/token")

# ── Row 3: API layer label ──
layer_label(3, "API & Authentication Layer", bg="065F46")

# ── Row 4: FastAPI Backend ──
api = full_merge(4)
arch_cell(api,
          title="FastAPI Backend",
          subtitle="src/backend/main.py  ·  Port 8000",
          pills=["RequestID Middleware", "CORS Middleware", "Lifespan Events", "ApiResponse[T] Envelope"],
          bg="059669", title_color="FFFFFF", subtitle_color="A7F3D0",
          pill_fg="064E3B")

# ── Row 5: Auth0 sub-box ──
auth_cell = full_merge(5)
set_cell_bg(auth_cell, "D1FAE5")
auth_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
p_auth = auth_cell.paragraphs[0]
p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_auth.paragraph_format.space_before = Pt(4)
p_auth.paragraph_format.space_after = Pt(4)
r_lock = p_auth.add_run("🔒  Auth0 JWT Guard  ")
r_lock.bold = True
r_lock.font.size = Pt(10)
r_lock.font.color.rgb = RGBColor.from_string("065F46")
r_pills = p_auth.add_run("RS256 Signature   ·   Audience / Issuer Check   ·   Expiry Validation   ·   JWKS Cache")
r_pills.font.size = Pt(9)
r_pills.italic = True
r_pills.font.color.rgb = RGBColor.from_string("047857")

# ── Row 6: arrow ──
arrow_row(diag, 6, label="run_agent( user_message, session_id )")

# ── Row 7: Agent layer label ──
layer_label(7, "AI Agent & Orchestration Layer", bg="6B21A8")

# ── Row 8: LangChain Agent ──
agent_cell = full_merge(8)
arch_cell(agent_cell,
          title="LangChain Agent Orchestrator",
          subtitle="src/backend/agent/job_agent.py",
          pills=["create_tool_calling_agent", "AgentExecutor", "RunnableWithMessageHistory", "Gemini (gemini-2.0-flash)"],
          bg="7C3AED", title_color="FFFFFF", subtitle_color="DDD6FE",
          pill_fg="4C1D95")

# ── Row 9: Memory | Langfuse ──
row9 = diag.rows[9]
mem_cell = row9.cells[0].merge(row9.cells[1])
lf_cell  = row9.cells[2]   # already separate after merge

arch_cell(mem_cell,
          title="Session Memory",
          subtitle="ChatMessageHistory per session_id",
          pills=["In-memory store", "Multi-turn context"],
          bg="A78BFA", title_color="FFFFFF", subtitle_color="EDE9FE", pill_fg="4C1D95")

arch_cell(lf_cell,
          title="Langfuse Tracer",
          subtitle="CallbackHandler",
          pills=["Token counts", "Latency", "Tool traces"],
          bg="8B5CF6", title_color="FFFFFF", subtitle_color="EDE9FE", pill_fg="4C1D95")

# ── Row 10: arrow ──
arrow_row(diag, 10, label="Tool Calls  (dynamic routing by the LLM)")

# ── Row 11: Tools layer label ──
layer_label(11, "Agent Tools Layer", bg="92400E")

# ── Row 12: Three tools ──
tools = [
    ("Job Search Tool",       "agent/tools/job_search.py",    ["search_jobs()", "SerpAPI client", "Markdown formatter"]),
    ("Resume Analyser Tool",  "agent/tools/resume_analyzer.py",["analyze_resume()", "Skills extraction", "Score / Gap analysis"]),
    ("Cover Letter Tool",     "agent/tools/cover_letter.py",   ["generate_cover_letter()", "4-para template", "300–400 words"]),
]
tool_bgs = ["B45309", "D97706", "F59E0B"]
for i, (title, sub, pills) in enumerate(tools):
    arch_cell(diag.rows[12].cells[i],
              title=title, subtitle=sub, pills=pills,
              bg=tool_bgs[i], title_color="FFFFFF", subtitle_color="FEF3C7", pill_fg="78350F")

# ── Row 13: arrows (3 cols, each own cell) ──
for cell in diag.rows[13].cells:
    set_cell_bg(cell, "F8FAFC")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    ar = p.add_run("↓")
    ar.font.size = Pt(14)
    ar.bold = True
    ar.font.color.rgb = RGBColor.from_string("2563EB")

# ── Row 14: External services label ──
layer_label(14, "External Services", bg="1E3A5F")

# ── Row 15: Three external services ──
ext_services = [
    ("SerpAPI  ·  Google Jobs",    "Real-time job listings",      ["Up to 10 results", "Filters: role + location"]),
    ("Google Gemini LLM",          "Resume Analysis  &  Gap Check",["gemini-2.0-flash", "OpenAI-compatible API"]),
    ("Google Gemini LLM",          "Cover Letter Generation",      ["gemini-2.0-flash", "OpenAI-compatible API"]),
]
ext_bgs = ["1E3A5F", "1E40AF", "1D4ED8"]
for i, (title, sub, pills) in enumerate(ext_services):
    arch_cell(diag.rows[15].cells[i],
              title=title, subtitle=sub, pills=pills,
              bg=ext_bgs[i], title_color="FFFFFF", subtitle_color="BFDBFE", pill_fg="DBEAFE")

# ── Row 16: Footer — other external services ──
footer_cell = full_merge(16)
set_cell_bg(footer_cell, "0F172A")
pf = footer_cell.paragraphs[0]
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf.paragraph_format.space_before = Pt(5)
pf.paragraph_format.space_after = Pt(5)
ext_icons = [
    ("Auth0", "JWT issuance & JWKS endpoint"),
    ("Langfuse Cloud", "Observability dashboard  ·  cloud.langfuse.com"),
    ("Vercel", "Serverless deployment  ·  Mangum ASGI adapter"),
]
for j, (svc, desc) in enumerate(ext_icons):
    if j:
        sep = pf.add_run("     |     ")
        sep.font.size = Pt(9)
        sep.font.color.rgb = RGBColor.from_string("4B5563")
    rs = pf.add_run(svc)
    rs.bold = True
    rs.font.size = Pt(9)
    rs.font.color.rgb = RGBColor.from_string("60A5FA")
    rd = pf.add_run(f"  {desc}")
    rd.font.size = Pt(8)
    rd.italic = True
    rd.font.color.rgb = RGBColor.from_string("9CA3AF")

doc.add_paragraph()

doc.add_page_break()

# ===========================================================================
# SECTION 6 — DATA FLOW
# ===========================================================================
heading(doc, "6. Data Flow — End-to-End Example", level=1)
add_horizontal_rule(doc)
body(doc, 'Example: User types "Find Python jobs in San Francisco and analyse my resume."')
doc.add_paragraph()

flow_steps = [
    ("1", "Frontend",  "User submits message + session_id. Streamlit calls POST /api/chat/public."),
    ("2", "Backend",   "RequestIDMiddleware attaches UUID. Router calls run_agent(message, session_id)."),
    ("3", "Agent",     "Retrieves ChatMessageHistory for session_id. Invokes Gemini LLM to plan steps."),
    ("4", "Tool Call", 'Agent decides to call search_jobs("Python Engineer", "San Francisco").'),
    ("5", "SerpAPI",   "Returns up to 10 job listings as JSON. Tool formats to markdown."),
    ("6", "Agent",     "Gemini receives tool result. Decides to also call analyze_resume(resume_text)."),
    ("7", "LLM Tool",  "Gemini analyses resume, returns skills, score, and improvement suggestions."),
    ("8", "Agent",     "Synthesises both tool results into a combined markdown response."),
    ("9", "Memory",    "Full exchange (human + AI messages) saved to ChatMessageHistory."),
    ("10", "Response", 'FastAPI wraps output in ApiResponse.ok(ChatResponse(response=..., session_id=...)).'),
    ("11", "Langfuse", "CallbackHandler submits trace: 2 tool calls, 3 LLM turns, total latency, tokens."),
    ("12", "Frontend", "Streamlit renders markdown response in chat bubble. User sees jobs + analysis."),
]

df_tbl = doc.add_table(rows=len(flow_steps) + 1, cols=3)
df_tbl.style = "Table Grid"
for i, hdr in enumerate(["Step", "Layer", "Action"]):
    c = df_tbl.rows[0].cells[i]
    set_cell_bg(c, "1E3A5F")
    r = c.paragraphs[0].add_run(hdr)
    r.bold = True; r.font.color.rgb = RGBColor(255, 255, 255); r.font.size = Pt(11)

for idx, (step, layer, action) in enumerate(flow_steps, start=1):
    row = df_tbl.rows[idx]
    bg = "EFF6FF" if idx % 2 == 0 else "FFFFFF"
    for c in row.cells:
        set_cell_bg(c, bg)
    row.cells[0].paragraphs[0].add_run(step).font.size = Pt(11)
    lr = row.cells[1].paragraphs[0].add_run(layer)
    lr.bold = True; lr.font.size = Pt(11); lr.font.color.rgb = RGBColor.from_string("2563EB")
    row.cells[2].paragraphs[0].add_run(action).font.size = Pt(11)

doc.add_page_break()

# ===========================================================================
# SECTION 7 — ERROR HANDLING
# ===========================================================================
heading(doc, "7. Error Handling Strategy", level=1)
add_horizontal_rule(doc)
body(doc, "All errors are mapped through a domain exception hierarchy and surfaced as "
         "structured ApiResponse envelopes with consistent HTTP status codes:")

err_tree = [
    ("AppError",             "500", "INTERNAL_SERVER_ERROR — base for all domain errors"),
    ("  ConfigurationError", "500", "Missing or invalid environment variable"),
    ("  Auth0Error",         "502", "Generic Auth0 failure"),
    ("    Auth0ConfigError", "500", "Auth0 not configured"),
    ("    Auth0CredentialsError", "401", "Invalid JWT or credentials"),
    ("    Auth0NetworkError","503", "Auth0 unreachable"),
    ("  GeminiError",        "502", "Generic LLM failure"),
    ("    GeminiConfigError","500", "API key missing"),
    ("    GeminiRateLimitError","429","Too many requests"),
    ("    GeminiQuotaExceededError","503","Quota exhausted"),
    ("    GeminiNetworkError","503","LLM unreachable"),
    ("    GeminiInvalidRequestError","400","Bad prompt/request"),
    ("  LangfuseError",      "502", "Generic observability failure (non-fatal)"),
    ("  SerpApiError",       "502", "Generic SerpAPI failure"),
    ("    SerpApiRateLimitError","429","SerpAPI rate limited"),
    ("    SerpApiNetworkError","503","SerpAPI unreachable"),
    ("  AgentError",         "502", "Agent execution failure"),
]
err_tbl = doc.add_table(rows=len(err_tree) + 1, cols=3)
err_tbl.style = "Table Grid"
for i, hdr in enumerate(["Exception Class", "HTTP", "Meaning"]):
    c = err_tbl.rows[0].cells[i]
    set_cell_bg(c, "1E3A5F")
    r = c.paragraphs[0].add_run(hdr)
    r.bold = True; r.font.color.rgb = RGBColor(255, 255, 255); r.font.size = Pt(10)

for idx, (cls, code, meaning) in enumerate(err_tree, start=1):
    row = err_tbl.rows[idx]
    bg = "EFF6FF" if idx % 2 == 0 else "FFFFFF"
    set_cell_bg(row.cells[0], bg); set_cell_bg(row.cells[2], bg)
    code_bg = {"400": "FEF3C7", "401": "FEE2E2", "429": "FEF3C7",
               "500": "FEE2E2", "502": "FFF7ED", "503": "FFF7ED"}.get(code, bg)
    set_cell_bg(row.cells[1], code_bg)
    cr = row.cells[0].paragraphs[0].add_run(cls.strip())
    cr.font.name = "Courier New"; cr.font.size = Pt(9)
    indent = len(cls) - len(cls.lstrip())
    row.cells[0].paragraphs[0].paragraph_format.left_indent = Pt(indent * 4)
    row.cells[1].paragraphs[0].add_run(code).font.size = Pt(10)
    row.cells[2].paragraphs[0].add_run(meaning).font.size = Pt(10)

doc.add_page_break()

# ===========================================================================
# SECTION 8 — TESTING
# ===========================================================================
heading(doc, "8. Testing Strategy", level=1)
add_horizontal_rule(doc)
body(doc, "The project includes 11 test modules (~1,689 lines) covering all layers:")

test_modules = [
    ("test_main.py",                 "FastAPI endpoint integration tests; validates request routing, response envelopes, error codes"),
    ("test_agent.py",                "Agent session management, tool routing, memory persistence"),
    ("test_auth.py",                 "Auth0 JWT RS256 validation — valid/expired/invalid tokens"),
    ("test_auth_token.py",           "Token acquisition via Resource Owner Password Grant"),
    ("test_llm.py",                  "LLM factory singleton, error translation for all 5 Gemini error types"),
    ("test_tools_job_search.py",     "SerpAPI integration, error handling, result formatting"),
    ("test_tools_resume_analyzer.py","Resume analysis with and without job description"),
    ("test_tools_cover_letter.py",   "Cover letter generation, validation edge cases"),
    ("test_schemas.py",              "Pydantic model validation — field constraints, custom validators"),
    ("test_langfuse.py",             "Langfuse handler creation, graceful degradation when credentials missing"),
    ("conftest.py",                  "Shared fixtures: mocked LLM, SerpAPI, Auth0, Langfuse, TestClient"),
]
test_tbl = doc.add_table(rows=len(test_modules) + 1, cols=2)
test_tbl.style = "Table Grid"
for i, hdr in enumerate(["Test Module", "Coverage Focus"]):
    c = test_tbl.rows[0].cells[i]
    set_cell_bg(c, "1E3A5F")
    r = c.paragraphs[0].add_run(hdr)
    r.bold = True; r.font.color.rgb = RGBColor(255, 255, 255); r.font.size = Pt(10)

for idx, (module, focus) in enumerate(test_modules, start=1):
    row = test_tbl.rows[idx]
    bg = "EFF6FF" if idx % 2 == 0 else "FFFFFF"
    set_cell_bg(row.cells[0], bg); set_cell_bg(row.cells[1], bg)
    mr = row.cells[0].paragraphs[0].add_run(module)
    mr.font.name = "Courier New"; mr.font.size = Pt(9); mr.font.color.rgb = RGBColor.from_string("7C3AED")
    row.cells[1].paragraphs[0].add_run(focus).font.size = Pt(10)

doc.add_paragraph()
body(doc, "Run the test suite:", bold=True)
code_block(doc, [
    "pip install -r requirements-test.txt",
    "pytest tests/ -v --cov=src/backend --cov-report=term-missing",
])

doc.add_page_break()

# ===========================================================================
# SECTION 9 — DEPLOYMENT
# ===========================================================================
heading(doc, "9. Deployment Configuration", level=1)
add_horizontal_rule(doc)

heading(doc, "Local Development", level=2, color="2563EB")
code_block(doc, [
    "# Backend",
    "cd src/backend",
    "pip install -r requirements.txt",
    "uvicorn main:app --reload --port 8000",
    "",
    "# Frontend",
    "cd src/frontend",
    "pip install -r requirements.txt",
    "streamlit run app.py",
])
doc.add_paragraph()

heading(doc, "Vercel Serverless Deployment", level=2, color="2563EB")
body(doc, "Both backend and frontend include vercel.json configuration:")
code_block(doc, [
    "# src/backend/vercel.json",
    '{',
    '  "builds": [{"src": "index.py", "use": "@vercel/python"}],',
    '  "routes": [{"src": "/(.*)", "dest": "index.py"}]',
    '}',
    "",
    "# Mangum ASGI adapter in index.py wraps the FastAPI app",
    "# for AWS Lambda / Vercel serverless function execution",
])

doc.add_page_break()

# ===========================================================================
# SECTION 10 — ENVIRONMENT VARIABLES
# ===========================================================================
heading(doc, "10. Environment Variables Reference", level=1)
add_horizontal_rule(doc)

env_vars = [
    ("GEMINI_API_KEY",       "Required", "Google API key for Gemini LLM"),
    ("GEMINI_MODEL_NAME",    "Optional", "Default: gemini-2.0-flash"),
    ("GEMINI_BASE_URL",      "Optional", "Default: Google OpenAI-compatible endpoint"),
    ("SERPAPI_API_KEY",      "Required", "SerpAPI key for Google Jobs search"),
    ("AUTH0_DOMAIN",         "Required", "Auth0 tenant domain (e.g. dev-xxx.auth0.com)"),
    ("AUTH0_CLIENT_ID",      "Required", "Auth0 application client ID"),
    ("AUTH0_CLIENT_SECRET",  "Required", "Auth0 application client secret"),
    ("AUTH0_AUDIENCE",       "Required", "Auth0 API audience identifier"),
    ("LANGFUSE_SECRET_KEY",  "Optional", "Langfuse secret key (observability disabled if missing)"),
    ("LANGFUSE_PUBLIC_KEY",  "Optional", "Langfuse public key"),
    ("LANGFUSE_HOST",        "Optional", "Default: https://cloud.langfuse.com"),
    ("BACKEND_URL",          "Optional", "Default: http://localhost:8000 (used by frontend)"),
]
env_tbl = doc.add_table(rows=len(env_vars) + 1, cols=3)
env_tbl.style = "Table Grid"
for i, hdr in enumerate(["Variable", "Required", "Description"]):
    c = env_tbl.rows[0].cells[i]
    set_cell_bg(c, "1E3A5F")
    r = c.paragraphs[0].add_run(hdr)
    r.bold = True; r.font.color.rgb = RGBColor(255, 255, 255); r.font.size = Pt(10)

for idx, (var, req, desc) in enumerate(env_vars, start=1):
    row = env_tbl.rows[idx]
    bg = "EFF6FF" if idx % 2 == 0 else "FFFFFF"
    set_cell_bg(row.cells[0], bg)
    req_bg = "DCFCE7" if req == "Required" else "FFF7ED"
    set_cell_bg(row.cells[1], req_bg); set_cell_bg(row.cells[2], bg)
    vr = row.cells[0].paragraphs[0].add_run(var)
    vr.font.name = "Courier New"; vr.font.size = Pt(9); vr.font.color.rgb = RGBColor.from_string("7C3AED")
    rr = row.cells[1].paragraphs[0].add_run(req)
    rr.font.size = Pt(10); rr.bold = True
    if req == "Required":
        rr.font.color.rgb = RGBColor.from_string("15803D")
    else:
        rr.font.color.rgb = RGBColor.from_string("B45309")
    row.cells[2].paragraphs[0].add_run(desc).font.size = Pt(10)

doc.add_paragraph()
add_horizontal_rule(doc)
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run(
    "Job Placement Agent — Sprint 11 Capstone Project  |  "
    "Build and Deploy Intelligent Conversational Agents  |  NIIT"
)
fr.font.size = Pt(9)
fr.font.color.rgb = RGBColor.from_string("6B7280")
fr.italic = True

# ===========================================================================
# Save
# ===========================================================================
out_path = "Job_Placement_Agent_Blueprint.docx"
doc.save(out_path)
print(f"Blueprint saved to: {out_path}")
